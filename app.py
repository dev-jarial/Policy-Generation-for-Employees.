import docx
import openai

# Set your OpenAI API key
openai.api_key = "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"

def generate_completion(api_key, conversation_history, max_tokens=7200):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {
                "role": "system",
                "content": "You are a helpful assistant."
            },
            *[
                {"role": "user", "content": msg}
                for msg in conversation_history
            ],
        ],
        max_tokens=max_tokens,
        n=1,
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()

def generate_document(conversation_history):
    document = docx.Document()
    completion = generate_completion(openai.api_key, conversation_history)
    document.add_paragraph("")
    for msg in conversation_history:
        document.add_paragraph(msg, style="List Bullet")
    document.add_paragraph("")
    document.add_paragraph(completion)
    document.save("prompt_completion.docx")

# Read the policy.txt file
with open("policy.txt", "r") as file:
    conversation_history = file.read().splitlines()

# Generate completion and save as a DOCX document
generate_document(conversation_history)
